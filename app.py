import sys
import os
import io
import json
import shutil
from dotenv import load_dotenv
import streamlit as st
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from markdown2 import markdown
from bs4 import BeautifulSoup

project_root = os.path.abspath(os.path.dirname(__file__))
sys.path.insert(0, project_root)
sys.path.insert(0, os.path.join(project_root, 'src'))

from src.crew import create_crew

load_dotenv()

st.set_page_config(page_title="Liquidity Risk Analysis Dashboard", layout="wide")
st.image("treasury_header.jpg", caption="Liquidity Risk Dashboard", width=400)
st.title("dY Liquidity Risk Analysis Dashboard")
st.markdown(
    """
    Upload your files and provide an objective to generate a professional liquidity risk report.
    """
)

objective = st.text_area(
    "Analysis Objective",
    value="Assess overall liquidity risk, focusing on cash forecasts and debt risks.",
    height=100,
    help="Guide the analyst (e.g., 'Prioritize covenant checks' or 'Forecast cash shortfalls')."
)

uploaded_files = st.file_uploader(
    "Upload Files (Excel recommended)",
    accept_multiple_files=True,
    help="Upload Excel, CSV, TSV, or JSON files."
)

if st.button("Analyse", key="analyse"):
    if not uploaded_files:
        st.error("Upload at least one file.")
    else:
        objective_text = objective.strip()
        if not objective_text:
            st.error("Provide an objective.")
        else:
            upload_dir = "uploads"
            os.makedirs(upload_dir, exist_ok=True)
            file_paths = []
            for uploaded_file in uploaded_files:
                file_path = os.path.join(upload_dir, uploaded_file.name)
                with open(file_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                file_paths.append(file_path)

            try:
                with st.spinner("Analyzing..."):
                    crew_result = create_crew(file_paths, objective_text)
                    if hasattr(crew_result, "raw"):
                        report_text = crew_result.raw or ""
                        if not report_text and getattr(crew_result, "json_dict", None):
                            report_text = json.dumps(crew_result.json_dict, indent=2)
                    elif isinstance(crew_result, str):
                        report_text = crew_result
                    else:
                        report_text = str(crew_result or "")

                    if not report_text.strip():
                        st.warning("Analysis completed but no readable report was produced.")
                    else:
                        st.markdown(report_text, unsafe_allow_html=True)
                        st.success("Analysis complete!")

                        html_report = markdown(report_text)
                        soup = BeautifulSoup(html_report, "html.parser")
                        text_content = soup.get_text()

                        doc = Document()
                        doc.add_heading("Liquidity Risk Analysis Report", 0)
                        for line in text_content.split('\n'):
                            line = line.strip()
                            if line:
                                doc.add_paragraph(line)
                        word_buffer = io.BytesIO()
                        doc.save(word_buffer)
                        word_buffer.seek(0)
                        st.download_button(
                            label="Download Word",
                            data=word_buffer,
                            file_name="liquidity_report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

                        pdf_buffer = io.BytesIO()
                        pdf = SimpleDocTemplate(pdf_buffer, pagesize=letter)
                        styles = getSampleStyleSheet()
                        story = [
                            Paragraph(line.strip(), styles["Normal"])
                            for line in text_content.split('\n')
                            if line.strip()
                        ]
                        pdf.build(story)
                        pdf_buffer.seek(0)
                        st.download_button(
                            label="Download PDF",
                            data=pdf_buffer,
                            file_name="liquidity_report.pdf",
                            mime="application/pdf"
                        )
            except Exception as exc:
                st.error(f"Error: {exc}")
            finally:
                if os.path.exists(upload_dir):
                    shutil.rmtree(upload_dir)
